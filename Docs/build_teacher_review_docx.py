from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"E:\match\LoongArch\Docs\teacher-review-one-flow-summary.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string("1F3A5F")
    p_fmt = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_fmt.append(shd)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(header)
        set_run_font(r, bold=True, color="0B2545")
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            set_run_font(r)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_width(table, widths)
    set_table_borders(table)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Number", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("厂区配电与动力设备检修知识助手 | 操作流程与功能说明")
    set_run_font(r, size=9, color="666666")


def build():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("厂区配电与动力设备检修知识助手")
    set_run_font(r, size=20, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("操作流程与功能说明")
    set_run_font(r, size=13, color="2E74B5")

    add_para(doc, "面向老师评审的简版说明：本文用一条完整操作流程说明系统有什么功能、能做什么，以及演示时如何体现赛题要求。")

    doc.add_heading("1. 系统能做什么", level=1)
    add_para(doc, "本系统面向制造企业厂区配电与动力设备检修场景，帮助一线检修人员快速完成从现场描述到作业卡生成、专家修正回流的完整闭环。")
    add_numbered(doc, [
        "描述现场故障现象。",
        "上传现场图片或红外图。",
        "自动检索检修知识、规程和历史案例。",
        "通过多 Agent 协同给出检修建议。",
        "展示故障原因、检测方法、处理措施之间的知识图谱。",
        "生成可执行的标准化检修作业卡。",
        "由专家审核和修正结果。",
        "将专家经验回流到知识库，后续同类问题优先引用。",
    ])
    add_para(doc, "一句话说明：工人把现场看到的问题告诉系统，系统给出有依据的检修建议和作业卡，专家修正后还能让系统下次回答得更准。")

    doc.add_heading("2. 一条完整操作流程", level=1)
    steps = [
        ("第一步：一线人员登录系统", "一线检修人员使用账号登录，进入工作台。工作台提供开始智能诊断入口、最近诊断记录和专家待审核数量提示。", "基础登录与角色区分；PC Web 可视化操作界面。"),
        ("第二步：描述现场故障", "用户进入智能诊断台，直接输入现场现象，例如“巡检发现 B 相电缆接头温度明显高于其他两相”。系统继续追问设备、型号、检修等级和是否有现场图片。", "智能接诊；不要求工人一开始就知道准确故障名称；适合新员工和现场快速使用。"),
        ("第三步：上传图片作为辅助信息", "用户上传现场图片、红外测温图或告警截图。系统识别疑似设备类型、异常部位、温度异常或可见故障痕迹，并要求用户确认不确定信息。", "多模态输入；图片辅助识别；人工确认避免误判。"),
        ("第四步：系统检索知识并给出建议", "系统根据用户描述、设备信息、图片识别结果和检修等级，自动检索知识库，输出可能故障、原因、检查方法、处理建议、安全注意事项和相似案例。", "RAG 知识检索；证据卡片；回答有依据，不是普通聊天式回答。"),
        ("第五步：多 Agent 协同会诊", "系统组织接诊 Agent、知识检索 Agent、领域专家 Agent、安全合规 Agent 共同分析，并汇总统一建议。", "多 Agent 协同；老师傅经验数字化；技术判断和安全合规同时参与。"),
        ("第六步：查看知识图谱", "系统展示当前故障相关知识图谱，如“电缆接头 -> 接头发热 -> 接触电阻增大 -> 红外测温 / 回路电阻测试 -> 停电检查与紧固 -> 安全要求”。", "知识图谱；故障原因链路可解释；展示检修知识关系。"),
        ("第七步：生成标准化作业卡", "用户点击生成作业卡。作业卡包含设备信息、故障现象、检修等级、风险等级、安全检查、操作步骤、验收标准、异常升级条件、知识来源和参与会诊的 Agent，并支持浏览器保存为 PDF。", "标准化作业指引；安全检查前置；输出可执行结果；支持 PDF 交付。"),
        ("第八步：专家审核与修正", "专家审核员登录系统，修正故障原因、检查方法、处理措施、作业卡步骤、知识图谱关系或 Agent 会诊意见，审核通过后进入知识库。", "专家审核；人工修正大模型输出；保证工业检修场景可靠性。"),
        ("第九步：知识回流，下一次回答更准", "专家修正通过后，系统将内容写入知识库、知识图谱和专家 Agent 经验记录。下一次遇到类似故障时，系统优先引用专家已审核经验，并显示“专家修正 · 已审核”。", "知识沉淀与更新；经验传承；系统越用越准。"),
    ]
    for title_text, body, value in steps:
        doc.add_heading(title_text, level=2)
        add_para(doc, body)
        add_para(doc, "体现功能：" + value, bold_lead="体现功能：")

    doc.add_heading("3. 演示时可以这样讲", level=1)
    add_para(doc, "本系统演示一条完整闭环：")
    add_code_block(doc, [
        "现场现象输入",
        "  -> 图片辅助识别",
        "  -> 知识库检索",
        "  -> 多 Agent 会诊",
        "  -> 知识图谱解释",
        "  -> 生成 PDF 作业卡",
        "  -> 专家审核修正",
        "  -> 知识回流",
    ])
    add_para(doc, "核心价值：")
    add_numbered(doc, [
        "一线人员不用翻大量手册。",
        "新员工不知道准确关键词也能发起诊断。",
        "每条建议都有知识来源。",
        "作业卡能直接指导检修。",
        "安全要求不会被遗漏。",
        "老师傅经验可以沉淀到系统中。",
        "系统部署在龙芯 + 银河麒麟环境，智能能力通过云端服务实现。",
    ])

    doc.add_heading("4. 首期演示范围", level=1)
    add_simple_table(
        doc,
        ["演示设备", "重点故障"],
        [
            ["电缆接头", "接头发热；绝缘老化/放电"],
            ["低压电机", "轴承过热/异响；绕组过热/绝缘异常"],
        ],
        [Inches(1.8), Inches(4.7)],
    )
    add_para(doc, "其他设备后续可通过配置知识库继续扩展。")

    doc.add_heading("5. 一句话总结", level=1)
    add_para(doc, "这是一个面向制造企业一线检修人员的智能检修助手：工人描述现场问题，系统通过知识库、多 Agent、知识图谱和专家回流机制，生成有依据、可执行、可持续更新的检修作业方案。")

    doc.core_properties.title = "厂区配电与动力设备检修知识助手：操作流程与功能说明"
    doc.core_properties.subject = "中国软件杯 A1 赛题项目说明"
    doc.core_properties.author = "项目组"
    doc.save(OUT)


if __name__ == "__main__":
    build()
